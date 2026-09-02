#!/usr/bin/env python3
"""Build the submission-ready Nightingale Care Note technical brief."""

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DIST = ROOT / "dist"
ARCHITECTURE_IMAGE = DOCS / "architecture.png"
OUTPUT = DIST / "Nightingale_Care_Note_Technical_Brief.docx"

INK = "17252A"
NAVY = "183D4E"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
PINE = "12635A"
PINE_LIGHT = "E7F2EF"
GRAY = "5C6970"
LIGHT_GRAY = "F2F4F7"
BORDER = "CDD5D9"
WHITE = "FFFFFF"
GOLD = "A66C1A"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    assert sum(widths) == PAGE_WIDTH_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def table_cell_text(cell, text, *, bold=False, color=INK, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def paragraph_border_bottom(paragraph, color=PINE, size=12, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, result, end])
    set_run_font(run, size=9, color=GRAY)


def configure_document(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DEEP_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # LibreOffice's macOS headless renderer applies odd/even header parts
    # inconsistently. Keep all variants intentionally empty so every rendered
    # page has the same clean margins and no clipped page furniture.
    for part in (
        section.header,
        section.first_page_header,
        section.even_page_header,
        section.footer,
        section.first_page_footer,
        section.even_page_footer,
    ):
        part.paragraphs[0].clear()


def add_real_bullet_definition(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_body(doc, text, *, bold_lead=None, after=6, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=size, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=size)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size)
    return p


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def draw_architecture(path):
    width, height = 1800, 560
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    try:
        regular = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 31)
        small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 24)
        bold = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 31)
    except OSError:
        regular = small = bold = ImageFont.load_default()

    boxes = [
        (45, 150, 330, 370, "Care Note\nworkspace", "browser"),
        (410, 150, 695, 370, "HTTP API", "signed session"),
        (775, 85, 1110, 260, "Policy + service", "RBAC • versions\nprovenance • review"),
        (775, 310, 1110, 485, "No-PHI gateway", "redact • hash\nlocal scribe"),
        (1190, 85, 1540, 260, "SQLite", "ACID records +\nimmutable versions"),
        (1190, 310, 1540, 485, "Approved model", "redacted text only"),
    ]
    for x1, y1, x2, y2, title, subtitle in boxes:
        is_trust_box = "PHI" not in title and "model" not in title
        fill = "#E7F2EF" if is_trust_box else "#FFF4DF"
        outline = PINE if is_trust_box else GOLD
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline=f"#{outline}", width=5)
        bbox = draw.multiline_textbbox((0, 0), title, font=bold, spacing=5, align="center")
        tx = (x1 + x2 - (bbox[2] - bbox[0])) / 2
        ty = y1 + 34
        draw.multiline_text((tx, ty), title, font=bold, fill=f"#{NAVY}", spacing=5, align="center")
        bbox2 = draw.multiline_textbbox((0, 0), subtitle, font=small, spacing=4, align="center")
        sx = (x1 + x2 - (bbox2[2] - bbox2[0])) / 2
        sy = y2 - (bbox2[3] - bbox2[1]) - 34
        draw.multiline_text((sx, sy), subtitle, font=small, fill=f"#{GRAY}", spacing=4, align="center")

    def arrow(start, end, color=PINE):
        draw.line((start, end), fill=f"#{color}", width=7)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 18, ey - 12), (ex - 18, ey + 12)], fill=f"#{color}")

    arrow((330, 260), (410, 260))
    arrow((695, 260), (775, 180))
    arrow((695, 285), (775, 395), GOLD)
    arrow((1110, 175), (1190, 175))
    arrow((1110, 395), (1190, 395), GOLD)
    draw.text((43, 31), "TRUST-CENTRIC REQUEST FLOW", font=regular, fill=f"#{GRAY}")
    draw.text((45, 505), "Production swap: OIDC + managed PostgreSQL/KMS; the policy boundary remains unchanged.", font=small, fill=f"#{GRAY}")
    image.save(path, format="PNG", optimize=True)


def add_first_page(doc, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("NIGHTINGALE 72HR BUILD")
    set_run_font(run, size=9.5, color=PINE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Nightingale Care Note")
    set_run_font(run, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("A trust-first longitudinal workspace for clinic collaboration")
    set_run_font(run, size=13.5, color=GRAY)

    for label, value in (
        ("Built by", "Qiufeng Wang"),
        ("Prototype", "Synthetic data only • Python + SQLite + browser-native UI"),
        ("Performance", "Authenticated Docker HTTP P95 4.721 ms (target ≤300 ms)"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=9.5, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=9.5, color=GRAY)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(3)
    paragraph_border_bottom(rule)

    add_heading(doc, "Product thesis", 1)
    add_body(
        doc,
        "The Care Note is a communication and trust layer beside the EHR - not another shared free-form document. Glance stays at three cards; Review is bounded at seven. A deterministic Consistency Watcher cites both immutable sources, while the Trust Passport requires evidence witnessing before any AI decision. Patient teach-back closes the understanding loop; access transparency and a tamper-evident audit chain make trust inspectable. Role-owned entries preserve authority and never silently merge clinician, staff, patient, or AI claims.",
        after=5,
    )

    add_heading(doc, "Architecture and request flow", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = p.add_run().add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.45))
    picture._inline.docPr.set("title", "Trust-centric request flow")
    picture._inline.docPr.set(
        "descr",
        "Care Note browser calls a signed HTTP API. Policy and service use SQLite, while AI text passes through a no-PHI gateway before an approved model.",
    )

    for item in (
        "Every API call resolves a short-lived signed actor, then enforces patient binding or clinic scope server-side.",
        "SQLite supplies ACID writes and optimistic locking in the prototype; the schema maps directly to PostgreSQL/RLS.",
        "AI text passes only through the visible redaction gateway; the Trust Passport retains source, verification, decision, bounded learning impact, and retention evidence.",
    ):
        add_bullet(doc, item, num_id)


def add_schema_table(doc):
    headers = ("Entity", "Core fields", "Trust role")
    rows = [
        ("Entry", "patient, role, type, visibility, section, current version", "Role-owned timeline unit"),
        ("EntryVersion", "entry, version, content, entities, risk, editor, time", "Append-only source snapshot"),
        ("Comment / Task", "entry, assignee, status, completion actor/time", "Auditable work without overwrite"),
        ("Highlight", "entry + version + offsets + quote + decision", "Exact span-level provenance"),
        ("Verification / conflict", "source version(s), verifier/decision, time", "Append-only evidence events"),
        ("AI note", "typed system entry + redaction digest/counts", "Distinct, reviewable authority"),
        ("ImportanceSignal", "clinic, feature, bounded weight, evidence", "Explainable local learning"),
        ("TeachBack", "instruction version, patient response, coverage, decision", "Human-confirmed understanding"),
        ("AuditLog", "metadata + prev_hash + event_hash", "Tamper-evident, content-free trace"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1700, 4350, 3310])
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], LIGHT_GRAY)
        table_cell_text(table.rows[0].cells[index], value, bold=True, color=NAVY, size=8.5)
    for entity, fields, trust in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell)
        table_cell_text(cells[0], entity, bold=True, color=DEEP_BLUE, size=8.2)
        table_cell_text(cells[1], fields, size=8.2)
        table_cell_text(cells[2], trust, size=8.2)
    return table


def add_second_page(doc, num_id):
    doc.add_page_break()
    add_heading(doc, "Data model and provenance", 1)
    add_body(
        doc,
        "A highlight is the trust anchor: entry_id + immutable entry_version + character offsets + quote. Resolution verifies the quote against that version. A short-lived evidence token binds actor, version, offsets, and quote digest before a decision is accepted. Verification is likewise an append-only event tied to one exact version, never a mutable timestamp. The same immutable versions power a Time Machine that reconstructs the exact note view and ten-second Glance at any past moment, honestly labelling learned priority and task state as current-only.",
        after=4,
        size=10.5,
    )
    add_schema_table(doc)

    add_heading(doc, "RBAC, privacy, and concurrency", 1)
    for item in (
        "Patients are bound to one patient ID and receive only patient-visible instructions; raw AI entries, internal comments, tasks, highlights, and audit data are denied even when a URL is guessed.",
        "Staff and clinicians edit only role-owned entries. Only clinicians make final AI/conflict/teach-back decisions after witnessing current-version evidence; task transitions require a clinic clinician or assigned/unassigned staff. Patients see content-free viewer metadata, never internal clinical content.",
        "Every edit carries expected_version inside a write transaction. Different sections save independently; competing writes yield one success and one deterministic HTTP 409 - never last-write-wins loss.",
    ):
        add_bullet(doc, item, num_id)

    add_heading(doc, "Consistency, review, and AI handling", 2)
    add_body(
        doc,
        "Consistency rules surface two immutable sources side by side and never diagnose. A source-linked pre-visit brief assembles safety, alerts, work, questions, and changes without creating new clinical conclusions. Review is capped at seven cards with a visible reason and no accept-all. Patient teach-back is bound to one instruction version; deterministic keyword coverage may flag a gap, but only a clinician confirms understanding. AI summaries remain typed system entries; the scribe preview shows raw browser memory versus the exact redacted payload before persistence, and no transcript leaves the machine.",
        after=3,
        size=10.5,
    )


def add_validation_table(doc):
    rows = [
        ("Automated suite", "33/33: RBAC; provenance; review; teach-back; audit chain; retention; time machine", "Pass"),
        ("Security tests", "Token/chain tamper; patient isolation; PHI preview; six safe local probes", "Pass"),
        ("Service benchmark", "300 reads after 20 warm-ups; median 6.249 ms; P95 6.622 ms", "≤300 ms"),
        ("HTTP benchmark", "300 Docker reads; signed session + RBAC + SQLite + JSON; median 3.575 ms; P95 4.721 ms", "≤300 ms"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1880, 5860, 1620])
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(("Evidence", "Result", "Status/target")):
        set_cell_shading(table.rows[0].cells[index], LIGHT_GRAY)
        table_cell_text(table.rows[0].cells[index], value, bold=True, color=NAVY, size=8.5)
    for evidence, result, status in rows:
        cells = table.add_row().cells
        table_cell_text(cells[0], evidence, bold=True, color=DEEP_BLUE, size=8.2)
        table_cell_text(cells[1], result, size=8.2)
        table_cell_text(cells[2], status, bold=True, color=PINE, size=8.2)
    return table


def add_third_page(doc, num_id):
    doc.add_page_break()
    add_heading(doc, "Learning, retention, and operational safety", 1)
    add_heading(doc, "Explainable importance", 2)
    add_body(
        doc,
        "Priority combines recency, explicit risk, unresolved work, clinical entities, and clinician authorship. Clinician interaction adjusts clinic-local feature weights, but total learned influence is clamped to ±4.0. The UI exposes used/remaining influence beside base score and rank movement. This is a ranking aid - not a diagnostic model.",
        size=10.5,
    )

    add_heading(doc, "Hybrid storage and data decay", 2)
    for item in (
        "Hot (≤90 days or safety protected): full content indexed for glance and timeline access.",
        "Warm (91-365 days): structured summary indexed; full immutable version retained.",
        "Cold (>365 days, low risk): encrypted archive; a provenance stub remains queryable.",
    ):
        add_bullet(doc, item, num_id)
    add_body(
        doc,
        "The clinician/admin Storage Lens exposes tier, age, protection, and policy for seeded hot/warm/cold history. High-risk, allergy, medication, and safety-net facts are never decayed solely because of age. Physical archival remains an explicit production integration.",
        after=3,
        size=10.2,
    )

    add_heading(doc, "Security boundary", 2)
    add_body(
        doc,
        "prepare_llm_payload removes known/labelled names, Singapore NRIC/FIN or labelled identifiers, and Singapore phone formats before any adapter can receive text. A no-persistence preview shows raw browser memory beside the exact redacted payload, counts, and digest. Content-free audit metadata is canonicalized into a per-clinic SHA-256 chain. A safe local sandbox demonstrates six blocked policy probes without scanning any external system. Production still requires managed ingress, PostgreSQL/object encryption, KMS keys, backups, retention, and legal hold.",
        size=10.2,
    )

    add_heading(doc, "Validation and measured performance", 1)
    add_validation_table(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Measurement scope: the HTTP figure includes local transport, signed-session verification, RBAC, SQLite, and JSON; browser rendering and production network still require deployed tracing/load tests.")
    set_run_font(run, size=8.5, color=GRAY, italic=True)

    add_heading(doc, "Assumptions and deliberate trade-offs", 2)
    add_body(
        doc,
        "The 72-hour build prioritizes the trust-critical collaboration path end to end. Deferred: production OIDC, live EHR/FHIR exchange, PostgreSQL RLS, CRDT rich text, terminology services, managed KMS/retention controls, and ambient voice diarization. These require clinical, security, and operational validation and are not represented as complete.",
        after=0,
        size=10.2,
    )


def audit_document(doc):
    section = doc.sections[0]
    assert round(section.left_margin.inches, 3) == 1.0
    assert round(section.right_margin.inches, 3) == 1.0
    assert round(section.top_margin.inches, 3) == 1.0
    assert round(section.bottom_margin.inches, 3) == 1.0
    assert round(section.header_distance.inches, 3) == 0.492
    assert round(section.footer_distance.inches, 3) == 0.492
    normal = doc.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size.pt == 11
    assert normal.paragraph_format.space_after.pt == 6
    assert abs(normal.paragraph_format.line_spacing - 1.10) < 0.001
    for table in doc.tables:
        widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid.findall(qn("w:gridCol"))]
        assert sum(widths) == PAGE_WIDTH_DXA
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == str(TABLE_INDENT_DXA)


def build():
    DOCS.mkdir(parents=True, exist_ok=True)
    DIST.mkdir(parents=True, exist_ok=True)
    draw_architecture(ARCHITECTURE_IMAGE)
    doc = Document()
    configure_document(doc)
    num_id = add_real_bullet_definition(doc)
    add_first_page(doc, num_id)
    add_second_page(doc, num_id)
    add_third_page(doc, num_id)
    audit_document(doc)
    doc.core_properties.title = "Nightingale Care Note - Technical Brief"
    doc.core_properties.subject = "Portfolio technical brief"
    doc.core_properties.author = "Qiufeng Wang"
    doc.core_properties.keywords = "Nightingale, care note, provenance, RBAC, clinical collaboration"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
